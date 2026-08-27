# -*- coding: utf-8 -*-
"""Native, controlled HTML->DOCX build of a SVAMP-style workforce one-pager.

Single source = the rendered report HTML (a `#page` div whose blocks use the
report's CSS conventions: `.title`, `.subtitle`, `h1`, `p`, `table.dem|live|
cmpgrid|trend`, `.xwrap`, `.emps`, `.footer`). Every element is mapped to a
native Word primitive — including the crosswalk, which is reconstructed from the
SVG as a native, fully-clickable table (a converge-to-target box for one SOC, a
program×SOC matrix for several) rather than rasterized. The high-DPI PNG remains
only as a fallback when the SVG can't be parsed. Produces an editable, faithful
.docx that survives the Google-Docs paste cleanly (native tables are the one
richly-supported primitive).

Usage:
    python3 build_docx.py [SRC.html] [OUT.docx] [CROSSWALK.png]

Defaults reproduce the "Manufacturing Technician" build. CROSSWALK.png is the
fallback raster (default /tmp/crosswalk.png) used only if the SVG crosswalk can't
be parsed into a native table; run shoot_xwalk_png.cjs first to have it on hand.

Deps: python-docx, beautifulsoup4.
"""
import os
import sys

import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from bs4 import BeautifulSoup

_RESEARCH = '/Users/dayonekoo/Desktop/code/kallipolis/research/swp-strategy'
SRC = sys.argv[1] if len(sys.argv) > 1 else f'{_RESEARCH}/svamp-pathway-49-9041-doc.html'
OUT = sys.argv[2] if len(sys.argv) > 2 else f'{_RESEARCH}/svamp-pathway-manufacturing-technician.docx'
XWALK = sys.argv[3] if len(sys.argv) > 3 else '/tmp/crosswalk.png'
FONT = 'Arial'
BYLINE_FONT = 'Days One'  # brand byline face (Google-native; substitutes in Word/Pages without it)

TEAL, BLUE, RED = '2a9d8f', '2e74b5', 'cc3333'
DARK, BODY, MUT = '2a3450', '33405a', '9099ab'
HFILL, TOTFILL, SECFILL = 'eef1f6', 'f2f6fc', 'eef1f6'
SOCCOL = {'lc1': TEAL, 'lc2': BLUE, 'lc3': RED, 'c1h': TEAL, 'c2h': BLUE, 'c3h': RED}

soup = BeautifulSoup(open(SRC, encoding='utf-8').read(), 'html.parser')
page = soup.find(id='page')
doc = Document()
st = doc.styles['Normal']
st.font.name = FONT
st.font.size = Pt(10)
st.font.color.rgb = RGBColor.from_string(BODY)
sec = doc.sections[0]
sec.page_width = Inches(8.5)
# DOCX_PAGE_H overrides page height (default US Letter) — set tall for single-image visual QA.
sec.page_height = Inches(float(os.environ.get('DOCX_PAGE_H', '11')))
sec.left_margin = sec.right_margin = Inches(0.7)
sec.top_margin = sec.bottom_margin = Inches(0.55)
CONTENT_W = 7.1


# ---------- low-level helpers ----------
def shade(cell, hexc):
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexc)
    cell._tc.get_or_add_tcPr().append(sh)


def left_accent(cell, hexc, sz=24):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    b = OxmlElement('w:left'); b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), '0'); b.set(qn('w:color'), hexc)
    borders.append(b); tcPr.append(borders)


def vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign'); va.set(qn('w:val'), 'center'); tcPr.append(va)


def grid(tbl, hexc='dfe3ea'):
    t = tbl._tbl
    el = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), hexc); el.append(e)
    t.tblPr.append(el)


def fill_width(tbl):
    """Pin the table full content width with fixed layout. python-docx leaves
    tblW=auto (renderers collapse narrow) AND leaves the gridCol widths equal even
    when cells set their own width via tcW — under fixed layout the gridCol wins, so
    a table like the crosswalk (4.0/0.5/2.6in cells) would render in equal thirds.
    Derive the column proportions from the first row's cell widths when present."""
    grid = tbl._tbl.find(qn('w:tblGrid'))
    if grid is None:
        return
    cols = grid.findall(qn('w:gridCol'))
    n = len(cols)
    total = int(round(CONTENT_W * 1440))
    rows = tbl._tbl.findall(qn('w:tr'))
    weights = None
    if rows:
        ws = []
        for tc in rows[0].findall(qn('w:tc')):
            tcPr = tc.find(qn('w:tcPr'))
            w = tcPr.find(qn('w:tcW')) if tcPr is not None else None
            span = tcPr.find(qn('w:gridSpan')) if tcPr is not None else None
            sp = int(span.get(qn('w:val'))) if span is not None else 1
            raw = w.get(qn('w:w')) if w is not None else None
            val = int(raw) if (raw and raw.isdigit()) else 0
            ws += [val / sp] * sp
        if len(ws) == n and sum(ws) > 0:
            weights = ws
    if weights is None:  # no per-cell widths → keep existing grid (already full) or equalize
        cur = [int(c.get(qn('w:w')) or 0) for c in cols]
        weights = cur if sum(cur) > 0 else [1] * n
    sw = sum(weights)
    widths = [int(total * w / sw) for w in weights]
    widths[-1] = total - sum(widths[:-1])
    for c, w in zip(cols, widths):
        c.set(qn('w:w'), str(w))
    tblPr = tbl._tbl.tblPr
    for tag in ('w:tblW', 'w:tblLayout'):
        e = tblPr.find(qn(tag))
        if e is not None:
            tblPr.remove(e)
    tw = OxmlElement('w:tblW'); tw.set(qn('w:type'), 'dxa'); tw.set(qn('w:w'), str(total))
    ly = OxmlElement('w:tblLayout'); ly.set(qn('w:type'), 'fixed')
    tblPr.append(tw); tblPr.append(ly)


def cellpad(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for k, v in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        e = OxmlElement('w:' + k); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa'); m.append(e)
    tcPr.append(m)


def run(p, text, size=10, bold=False, color=BODY, italic=False, font=FONT):
    r = p.add_run(text); r.font.name = font; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic; r.font.color.rgb = RGBColor.from_string(color)
    return r


def hyperlink(p, url, text, color=BLUE, size=9, font=FONT):
    r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement('w:hyperlink'); link.set(qn('r:id'), r_id)
    rr = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    for tag, val in (('w:color', color),):
        e = OxmlElement(tag); e.set(qn('w:val'), val); rPr.append(e)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'), font); rf.set(qn('w:hAnsi'), font); rPr.append(rf)
    rr.append(rPr); t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text; rr.append(t)
    link.append(rr); p._p.append(link)


def para(space_before=2, space_after=4):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(space_before); pf.space_after = Pt(space_after); pf.line_spacing = 1.12
    return p


def runs_from(el, p, size=10, color=BODY, font=FONT):
    """emit a <p>'s inline content as runs, honoring <b> and <a>."""
    for node in el.children:
        nm = getattr(node, 'name', None)
        if nm is None:
            txt = str(node).replace('\xa0', ' ')
            if txt.strip():
                run(p, txt, size=size, color=color, font=font)
            elif txt:  # whitespace-only node between inline elements → keep one space
                run(p, ' ', size=size, color=color, font=font)
        elif nm == 'b':
            run(p, node.get_text(' ', strip=True), size=size, bold=True, color=DARK, font=font)
        elif nm == 'i':
            run(p, node.get_text(' ', strip=True), size=size, italic=True, color=color, font=font)
        elif nm == 'a':
            hyperlink(p, node.get('href', ''), node.get_text(' ', strip=True), size=size, font=font)
        else:
            run(p, node.get_text(' ', strip=True), size=size, color=color, font=font)


def rows_of(table):
    out = []
    for tr in table.find_all('tr'):
        cells = []
        for c in tr.find_all(['th', 'td'], recursive=False):
            cells.append({'text': c.get_text(' ', strip=True).replace('\xa0', ' '), 'th': c.name == 'th',
                          'colspan': int(c.get('colspan', 1)), 'rowspan': int(c.get('rowspan', 1)),
                          'cls': c.get('class', []), 'el': c})
        if cells:
            out.append({'cells': cells, 'cls': tr.get('class', [])})
    return out


# ---------- block builders ----------
def add_title(text):
    text = text.replace('\xa0', '').strip()
    # keep the title on a single line: shrink to fit the content width (cap 20, floor 13)
    size = max(13.0, min(20.0, CONTENT_W * 72.0 / (len(text) * 0.58)))
    p = para(0, 1); p.paragraph_format.space_after = Pt(1)
    run(p, text, size=size, bold=True, color=DARK)
    # accent rule — short, tight blank line carrying the bottom border so the title sits close to it
    hr = para(0, 6); hr.paragraph_format.line_spacing = 1.0
    pPr = hr._p.get_or_add_pPr(); b = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom'); bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '10')
    bot.set(qn('w:space'), '2'); bot.set(qn('w:color'), BLUE); b.append(bot); pPr.append(b)


def add_lede(text):
    p = para(0, 8); runs = text.strip()
    run(p, runs, size=10.5, color=BODY)


def add_heading(text):
    p = para(8, 3); run(p, text.replace('\xa0', '').strip(), size=13, bold=True, color=BLUE)


def std_table(rows, widths=None, header=True, num_from=2, totalcls='tot'):
    ncol = max(sum(c['colspan'] for c in r['cells']) for r in rows)
    tbl = doc.add_table(rows=0, cols=ncol); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid(tbl)
    for r in rows:
        cells = tbl.add_row().cells
        ci = 0
        istot = totalcls in r['cls']
        ishdr = all(c['th'] for c in r['cells']) and r is rows[0]
        for c in r['cells']:
            cell = cells[ci]
            for j in range(1, c['colspan']):
                cell.merge(cells[ci + j])
            p = cell.paragraphs[0]; p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
            isnum = ci >= num_from
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if isnum else WD_ALIGN_PARAGRAPH.LEFT
            col = MUT if c['text'] in ('—', '-') else (DARK if (ishdr or istot) else BODY)
            run(p, c['text'], size=8.5, bold=(ishdr or istot), color=col)
            if ishdr:
                shade(cell, HFILL)
            if istot:
                shade(cell, TOTFILL)
            cellpad(cell)
            ci += c['colspan']
    return tbl


def add_trend(table):
    rows = rows_of(table)
    tbl = doc.add_table(rows=0, cols=6); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; grid(tbl)
    for ri, r in enumerate(rows):
        cells = tbl.add_row().cells
        istot = 'tot' in r['cls']; ishdr = ri == 0
        # Credential-mix sub-row under the member college. Without this branch the
        # <b> in cell 0 renders bold+dark exactly like a college name, so the docx
        # keeps the rows but loses the hierarchy that makes them read as a
        # breakdown — the flatten-not-drop failure the link-parity gate can't see.
        istier = 'tier' in r['cls']
        for ci, c in enumerate(r['cells']):
            cell = cells[ci]; p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
            if ci == 0 and istier:
                p.paragraph_format.left_indent = Pt(12)
                b = c['el'].find('b')
                run(p, (b.get_text(strip=True) if b else c['text']), size=8, bold=False, color=MUT)
            elif ci == 0 and not ishdr and not istot:
                b = c['el'].find('b'); span = c['el'].find('span')
                run(p, (b.get_text(strip=True) if b else c['text']), size=9, bold=True, color=DARK)
                if span:
                    p2 = cell.add_paragraph(); p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(0)
                    run(p2, span.get_text(' ', strip=True), size=7.5, color=MUT)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
                col = MUT if c['text'] in ('—', '-') else (DARK if (ishdr or istot) else BODY)
                run(p, c['text'], size=8 if istier else 8.5,
                    bold=(ishdr or istot), color=MUT if istier else col)
            if ishdr:
                shade(cell, HFILL)
            if istot:
                shade(cell, TOTFILL)
            cellpad(cell)


def add_live(table):
    # The occupation column rowspans every posting for a SOC (single-SOC reports
    # group all postings under one occupation cell), so later rows carry only
    # [employer, title-with-link]. Resolve the grid honoring rowspan and
    # vertical-merge the occupation cell, else those rows shift left and the
    # posting links land in a column the walker never reads.
    rows = rows_of(table)
    ncol = sum(c['colspan'] for c in rows[0]['cells'])
    tbl = doc.add_table(rows=len(rows), cols=ncol); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; grid(tbl)
    span_left = [0] * ncol  # rows below still covered by a rowspan in this column
    for ri, r in enumerate(rows):
        ishdr = ri == 0
        accent = next((SOCCOL[c] for c in r['cls'] if c in SOCCOL), None)
        ci = 0
        for c in r['cells']:
            while ci < ncol and span_left[ci] > 0:  # covered from above → extend the merge
                tbl.cell(ri - 1, ci).merge(tbl.cell(ri, ci)); span_left[ci] -= 1; ci += 1
            if ci >= ncol:
                break
            cell = tbl.cell(ri, ci); p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
            cls = c['cls']
            if ishdr:
                run(p, c['text'], size=8, bold=True, color='5a6577'); shade(cell, HFILL)
            elif 'lsoc' in cls:  # occupation + SOC sub (rowspans its postings)
                txt = c['el'].get_text('\n', strip=True).split('\n')
                run(p, txt[0], size=9, bold=True, color=DARK)
                if len(txt) > 1:
                    p2 = cell.add_paragraph(); p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(0)
                    run(p2, txt[1], size=7.5, color=MUT)
                if accent:
                    left_accent(cell, accent)
            elif 'lemp' in cls:  # employer
                run(p, c['text'], size=9, bold=True, color=DARK)
            else:  # ltit — the linked posting title
                a = c['el'].find('a')
                if a:
                    hyperlink(p, a.get('href', ''), a.get_text(' ', strip=True).replace('↗', '').strip() + '  ↗', size=9)
                else:
                    run(p, c['text'], size=9, color=BODY)
            cellpad(cell)
            if c['rowspan'] > 1:
                span_left[ci] = c['rowspan'] - 1
            ci += c['colspan']


def add_cmpgrid(table):
    # Two layouts share .cmpgrid: multi-SOC (cols = occupations, sec rows are
    # full-width KSA dividers) and single-SOC (cols = Knowledge/Skills/Abilities/
    # Technology, a sec row carries per-column sub-headers). Derive the column
    # count from the widest row and merge by colspan so both render.
    rows = rows_of(table)
    ncol = max(sum(c['colspan'] for c in r['cells']) for r in rows)
    tbl = doc.add_table(rows=0, cols=ncol); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; grid(tbl)
    for ri, r in enumerate(rows):
        cells = tbl.add_row().cells
        issec = 'sec' in r['cls']; isdesc = 'descrow' in r['cls']; ishdr = ri == 0
        ci = 0
        for c in r['cells']:
            cell = cells[ci]
            for j in range(1, c['colspan']):
                cell.merge(cells[ci + j])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
            if ishdr:  # colored occupation header(s) — one per SOC, or one full-width
                hexc = next((SOCCOL[x] for x in c['cls'] if x in SOCCOL), HFILL)
                lines = c['el'].get_text('\n', strip=True).split('\n')
                run(p, lines[0], size=9, bold=True, color='ffffff')
                if len(lines) > 1:
                    p2 = cell.add_paragraph(); p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(0)
                    run(p2, lines[1], size=7, color='ffffff')
                shade(cell, hexc)
            elif issec:  # section band — full-width divider OR per-column K/S/A/T sub-headers
                run(p, c['text'].upper(), size=8, bold=True, color=DARK); shade(cell, SECFILL)
            elif isdesc:  # description + its O*NET Occupation Summary link
                link = c['el'].find('a')
                desc = c['el'].get_text(' ', strip=True)
                lt = link.get_text(' ', strip=True) if link else ''
                if lt and lt in desc:
                    desc = desc[:desc.rfind(lt)].strip()
                run(p, desc, size=8.5, color=DARK, italic=True)
                if link:
                    p2 = cell.add_paragraph(); p2.paragraph_format.space_before = Pt(5); p2.paragraph_format.space_after = Pt(0)
                    hyperlink(p2, link.get('href', ''), lt, size=9)
            else:
                run(p, c['text'], size=8.5, color=BODY)
            cellpad(cell)
            ci += c['colspan']


def add_emps(div):
    p = para(2, 4)
    for node in div.children:
        nm = getattr(node, 'name', None)
        if nm == 'br':
            p = para(1, 2); continue
        if nm is None:
            t = str(node).replace('\xa0', ' ')
            if t.strip():
                run(p, t, size=10, color=BODY)
        elif nm == 'span' and 'e' in node.get('class', []):
            run(p, node.get_text(' ', strip=True), size=10, bold=True, color=DARK)
        elif nm == 'a':
            hyperlink(p, node.get('href', ''), node.get_text(' ', strip=True), size=9)
        else:
            run(p, node.get_text(' ', strip=True), size=10, color=BODY)


def add_footer(div):
    para(6, 2)
    for line in div.get_text('\n', strip=True).split('\n'):
        if not line.strip():
            continue
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(1)
        run(p, line.strip(), size=7.5, color=MUT, italic=True)


def add_image():
    if not os.path.exists(XWALK):
        return  # no rasterized crosswalk on hand → rely on add_xwalk_legend's link caption
    p = para(4, 4); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(XWALK, width=Inches(CONTENT_W))


def add_xwalk_legend(div):
    """Fallback only (used when the crosswalk renders as a rasterized funnel PNG):
    one quiet centered line of the per-college program links, which are otherwise
    lost to rasterization. The native crosswalk table carries them inline instead."""
    links = div.find_all('a')
    if not links:
        return
    p = para(1, 8); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, 'Program pages:  ', size=8, color=MUT, italic=True)
    for i, a in enumerate(links):
        coll = a.find_previous('text')
        name = coll.get_text(' ', strip=True) if coll else a.get_text(' ', strip=True)
        if i:
            run(p, '    ·    ', size=8, color=MUT)
        hyperlink(p, a.get('href', ''), name, size=8)


def _xwalk_data(div):
    """Recover the crosswalk graph from the rendered SVG: program nodes (left,
    344-wide rects), SOC occupation nodes (right, 254-wide rects), and the
    program→SOC edges (matched off each bezier's start/end y)."""
    import re as _re
    sv = div.find('svg') or div
    progy = sorted(float(r.get('y', 0)) for r in sv.find_all('rect') if abs(float(r.get('width', 0)) - 344) < 2)
    occr = sorted(((float(r.get('y', 0)), float(r.get('height', 0))) for r in sv.find_all('rect')
                   if abs(float(r.get('width', 0)) - 254) < 2), key=lambda t: t[0])
    progs = []
    for i, a in enumerate(sv.find_all('a')):
        coll = a.find_previous('text')
        # the TOP line is the <text> AFTER the program's own <text> — find_next from the
        # link itself would descend into that inner <text> and return the program name.
        prog_t = a.find('text')
        topn = prog_t.find_next('text') if prog_t is not None else a.find_next('text')
        progs.append({'college': coll.get_text(' ', strip=True) if coll else '',
                      'program': a.get_text(' ', strip=True), 'href': a.get('href', ''),
                      'top': topn.get_text(' ', strip=True) if topn else '',
                      'mid': (progy[i] + 30) if i < len(progy) else 0})
    texts = [t.get_text(' ', strip=True) for t in sv.find_all('text')]
    tops = [i for i, t in enumerate(texts) if t.startswith('TOP ')]
    occs, buf = [], []
    for t in (texts[max(tops) + 1:] if tops else texts):
        if t.startswith('SOC '):
            occs.append({'title': ' '.join(buf), 'soc': t}); buf = []
        else:
            buf.append(t)
    for i, o in enumerate(occs):
        o['mid'] = (occr[i][0] + occr[i][1] / 2) if i < len(occr) else 0
    near = lambda v, arr: min(range(len(arr)), key=lambda i: abs(arr[i]['mid'] - v)) if arr else -1
    edges = set()
    for path in sv.find_all('path'):
        nums = [float(n) for n in _re.findall(r'-?\d+\.?\d*', path.get('d', ''))]
        if len(nums) >= 8:
            edges.add((near(nums[1], progs), near(nums[-1], occs)))
    return progs, occs, edges


def _xwalk_progcell(cell, p, w, pad=48):
    cell.width = w
    par = cell.paragraphs[0]; par.paragraph_format.space_before = Pt(2); par.paragraph_format.space_after = Pt(1)
    run(par, p['college'], size=9, bold=True, color=DARK)
    p2 = cell.add_paragraph(); p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(1)
    hyperlink(p2, p['href'], p['program'], size=9)
    if p['top']:
        p3 = cell.add_paragraph(); p3.paragraph_format.space_before = Pt(0); p3.paragraph_format.space_after = Pt(2)
        run(p3, p['top'], size=7.5, color=MUT)
    cellpad(cell, top=pad, bottom=pad, left=110, right=110)


def add_xwalk_table(div):
    """Fully-native, fully-clickable crosswalk — no rasterization, so program-name
    links survive AND the figure survives a Google-Docs paste (a table is the one
    richly-supported primitive). One SOC → programs converge (arrows) to the target
    box. Multiple SOCs → a crosswalk matrix: programs × SOC columns, a colored dot
    at each program→SOC edge."""
    progs, occs, edges = _xwalk_data(div)
    if not progs or not occs:
        return False
    accents = [TEAL, BLUE, RED]
    if len(occs) == 1:  # converge-to-one-target
        n = len(progs)
        tbl = doc.add_table(rows=n, cols=3); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; tbl.autofit = False
        for ri, p in enumerate(progs):
            _xwalk_progcell(tbl.cell(ri, 0), p, Inches(4.0))
            left_accent(tbl.cell(ri, 0), accents[ri % len(accents)])
            c1 = tbl.cell(ri, 1); c1.width = Inches(0.5); vcenter(c1)
            pa = c1.paragraphs[0]; pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run(pa, '→', size=13, color=MUT)
        tgt = tbl.cell(0, 2).merge(tbl.cell(n - 1, 2)); tgt.width = Inches(2.6)
        shade(tgt, TEAL); vcenter(tgt); cellpad(tgt, top=80, bottom=80, left=100, right=100)
        tp = tgt.paragraphs[0]; tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(tp, occs[0]['title'], size=11, bold=True, color='ffffff')
        tp2 = tgt.add_paragraph(); tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(tp2, occs[0]['soc'], size=8, color='ffffff')
        return True
    # crosswalk matrix: rows = programs, columns = SOCs, dot at each edge
    ncol = 1 + len(occs)
    tbl = doc.add_table(rows=1 + len(progs), cols=ncol); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False; grid(tbl)
    W0 = Inches(3.3); WS = Inches((CONTENT_W - 3.3) / len(occs))
    # header: blank label cell + colored SOC columns
    h0 = tbl.cell(0, 0); h0.width = W0; shade(h0, HFILL); cellpad(h0)
    run(h0.paragraphs[0], 'College program', size=8, bold=True, color='5a6577')
    for si, o in enumerate(occs):
        c = tbl.cell(0, si + 1); c.width = WS; shade(c, accents[si % len(accents)]); vcenter(c)
        cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(cp, o['title'], size=8, bold=True, color='ffffff')
        cp2 = c.add_paragraph(); cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(cp2, o['soc'], size=7, color='ffffff')
        cellpad(c)
    # program rows
    for pi, p in enumerate(progs):
        _xwalk_progcell(tbl.cell(pi + 1, 0), p, W0, pad=28)
        for si in range(len(occs)):
            c = tbl.cell(pi + 1, si + 1); c.width = WS; vcenter(c)
            cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if (pi, si) in edges:
                run(cp, '●', size=12, color=accents[si % len(accents)])
            else:
                run(cp, '·', size=12, color='c8cdd8')
            cellpad(c)
    return True


# ---------- ordered walk ----------
def emit(el):
    nm = getattr(el, 'name', None)
    if nm is None:
        return
    cls = el.get('class', [])
    if 'title' in cls:
        add_title(el.get_text(' ', strip=True))
    elif 'subtitle' in cls:
        t = el.get_text(' ', strip=True)
        if t:
            add_lede(t)
    elif nm == 'h1':
        add_heading(el.get_text(' ', strip=True))
    elif nm == 'p':
        t = el.get_text(' ', strip=True)
        if not t:
            return
        if 'tnote' in cls:
            p = para(1, 4); run(p, t, size=8.5, color=MUT, italic=True)
        elif 'tnar' in cls:
            p = para(6, 2); runs_from(el, p, size=10, color='46536b')
        elif 'srcdash' in cls:
            p = para(2, 4); runs_from(el, p, size=10)
        elif 'srcsec' in cls:
            p = para(8, 3); runs_from(el, p, size=10)
        else:
            p = para(2, 5); runs_from(el, p, size=10.5)
    elif nm == 'table':
        if 'dem' in cls:
            std_table(rows_of(el))
        elif 'live' in cls:
            add_live(el)
        elif 'cmpgrid' in cls:
            add_cmpgrid(el)
        elif 'trend' in cls:
            add_trend(el)
    elif 'xwrap' in cls:
        # Native crosswalk (clickable, paste-safe) by default; fall back to the
        # rasterized funnel PNG + a link caption if the SVG can't be parsed.
        if not add_xwalk_table(el):
            add_image()
            add_xwalk_legend(el)
    elif 'emps' in cls:
        add_emps(el)
    elif 'footer' in cls:
        add_footer(el)
    elif 'byline' in cls:
        # left-aligned, just under the title rule, in the brand byline face
        p = para(0, 10)
        runs_from(el, p, size=10, color=DARK, font=BYLINE_FONT)
    elif 'srclist' in cls:
        for item in el.find_all('div', recursive=False):
            p = para(0, 1); p.paragraph_format.left_indent = Inches(0.16)
            runs_from(item, p, size=10)
    elif 'demstat' in cls:
        pass
    else:
        for ch in el.children:
            emit(ch)


for child in page.children:
    emit(child)

for _t in doc.tables:  # full-bleed every table
    fill_width(_t)

doc.save(OUT)
print('saved', OUT)
print('paragraphs:', len(doc.paragraphs), '| tables:', len(doc.tables))
